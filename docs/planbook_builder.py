from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


OUT = Path(__file__).with_name("计划书_优化扩充版.docx")

# A4 Chinese competition-plan override based on the narrative_proposal preset.
# The override keeps the proposal readable on A4, uses PingFang SC for CJK text,
# and keeps tables and lists explicit rather than relying on Word defaults.
TABLE_W = 9360
TABLE_IND = 120

FONT_BODY = "PingFang SC"
FONT_HEAD = "Heiti SC"
NAVY = "173B57"
BLUE = "2F6F8F"
TEAL = "2C8C85"
INK = "26343D"
MUTED = "667480"
LIGHT_BLUE = "EEF4F7"
LIGHT_TEAL = "EAF6F3"
LIGHT_GOLD = "FFF7E7"
BORDER = "D7E1E6"
WHITE = "FFFFFF"
GOLD = "B8832F"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_run_font(run, name=FONT_BODY, size=10.5, color=INK, bold=None, italic=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def set_style_font(style, name=FONT_BODY, size=10.5, color=INK, bold=None):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    if bold is not None:
        style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV", "start", "end"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key in ("sz", "val", "color", "space"):
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_table_geometry(table, widths, indent=TABLE_IND):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_w)
            tc_w.set(qn("w:w"), str(widths[min(idx, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def set_paragraph_border(paragraph, edge="bottom", color=BORDER, size="8", space="1"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color)


def add_field(paragraph, code, size=8.5):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {code} "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    set_run_font(run, size=size, color=MUTED)


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def set_cell_text(cell, text, size=9.2, color=INK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    for index, line in enumerate(str(text).split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=size, color=color, bold=bold)
    return paragraph


def set_table_style(table, header=True, header_fill=LIGHT_BLUE, body_size=9.2):
    for ridx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_shading(cell, header_fill if header and ridx == 0 else WHITE)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "5", "color": BORDER},
                bottom={"val": "single", "sz": "5", "color": BORDER},
                start={"val": "single", "sz": "5", "color": BORDER},
                end={"val": "single", "sz": "5", "color": BORDER},
            )
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    set_run_font(run, size=9.0 if ridx == 0 else body_size, color=NAVY if ridx == 0 else INK, bold=(ridx == 0))
        if header and ridx == 0:
            repeat_header(row)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, body_size=9.2, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        align = alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_text(table.rows[0].cells[idx], header, size=9.0, color=NAVY, bold=True, align=align)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            align = alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], value, size=body_size, color=INK, align=align)
    set_table_style(table, header=True, header_fill=header_fill, body_size=body_size)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(1)
    spacer.paragraph_format.space_after = Pt(1)
    return table


def add_callout(doc, label, text, fill=LIGHT_TEAL, accent=TEAL, size=10.0):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_W])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(
        cell,
        top={"val": "single", "sz": "5", "color": fill},
        bottom={"val": "single", "sz": "5", "color": fill},
        start={"val": "single", "sz": "22", "color": accent},
        end={"val": "single", "sz": "5", "color": fill},
    )
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.12
    r1 = p.add_run(label + "  ")
    set_run_font(r1, size=size, color=accent, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=size, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    return table


def add_body(doc, text, after=5.5, before=0, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=10.3, color=INK):
    p = doc.add_paragraph(style="Normal")
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    if first_line and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        p.paragraph_format.first_line_indent = Pt(20)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color)
    return p


def add_bullet(doc, text, size=10.0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Mm(6.5)
    p.paragraph_format.first_line_indent = Mm(-4.2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run_font(run, size=size, color=INK)
    return p


def new_numbering_instance(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum")) if node.get(qn("w:abstractNumId"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num")) if node.get(qn("w:numId"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "700")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "700")
    ind.set(qn("w:hanging"), "460")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def set_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def add_numbered(doc, text, size=10.0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Mm(7.0)
    p.paragraph_format.first_line_indent = Mm(-4.6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.12
    previous = doc.paragraphs[-2] if len(doc.paragraphs) >= 2 else None
    if previous is None or previous.style.name != "List Number":
        doc._xinyu_active_num_id = new_numbering_instance(doc)
    set_numbering(p, doc._xinyu_active_num_id)
    run = p.add_run(text)
    set_run_font(run, size=size, color=INK)
    return p


def add_section_heading(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    p.paragraph_format.keep_with_next = True
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph(text, style="Heading 2")
    p.paragraph_format.keep_with_next = True
    return p


def add_minor_heading(doc, text):
    p = doc.add_paragraph(text, style="Heading 3")
    p.paragraph_format.keep_with_next = True
    return p


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(11)
    r = p.add_run(text.upper())
    set_run_font(r, name=FONT_HEAD, size=10, color=TEAL, bold=True)
    return p


def add_toc_entry(doc, number, title, page):
    p = doc.add_paragraph(style="TOC Entry")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.tab_stops.add_tab_stop(Mm(164), WD_TAB_LEADER.DOTS)
    r = p.add_run(f"{number}  {title}")
    set_run_font(r, size=10.4, color=INK)
    r2 = p.add_run(f"\t{page}")
    set_run_font(r2, size=10.4, color=MUTED)
    return p


def add_meta_row(table, row_index, label, value):
    cells = table.rows[row_index].cells
    set_cell_text(cells[0], label, size=9.2, color=TEAL, bold=True)
    set_cell_text(cells[1], value, size=9.2, color=INK)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(17)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, FONT_BODY, 10.3, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.line_spacing = 1.18

    h1 = styles["Heading 1"]
    set_style_font(h1, FONT_HEAD, 16, NAVY, bold=True)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    set_style_font(h2, FONT_HEAD, 11.6, TEAL, bold=True)
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(3)
    h2.paragraph_format.line_spacing = 1.0
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    set_style_font(h3, FONT_HEAD, 10.4, BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(4)
    h3.paragraph_format.space_after = Pt(2)
    h3.paragraph_format.line_spacing = 1.0
    h3.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        s = styles[style_name]
        set_style_font(s, FONT_BODY, 10.0, INK)
        s.paragraph_format.space_before = Pt(0)
        s.paragraph_format.space_after = Pt(2.5)
        s.paragraph_format.line_spacing = 1.12

    for name in ("TOC Entry", "Cover Metadata", "Small Note"):
        if name not in styles:
            styles.add_style(name, 1)
    set_style_font(styles["TOC Entry"], FONT_BODY, 10.4, INK)
    set_style_font(styles["Cover Metadata"], FONT_BODY, 9.2, INK)
    set_style_font(styles["Small Note"], FONT_BODY, 8.6, MUTED)

    header = section.header
    hp = header.paragraphs[0]
    clear_paragraph(hp)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(3)
    left = hp.add_run("心语 | 项目计划书")
    set_run_font(left, name=FONT_HEAD, size=8.4, color=MUTED, bold=True)
    right = hp.add_run("\t东南大学成贤学院大学生创新大赛")
    set_run_font(right, size=8.4, color=MUTED)
    hp.paragraph_format.tab_stops.add_tab_stop(Mm(170))
    set_paragraph_border(hp, edge="bottom", color=BORDER, size="5", space="2")

    first_header = section.first_page_header
    clear_paragraph(first_header.paragraphs[0])
    first_footer = section.first_page_footer
    clear_paragraph(first_footer.paragraphs[0])

    footer = section.footer
    fp = footer.paragraphs[0]
    clear_paragraph(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a = fp.add_run("心语科创团队  ·  ")
    set_run_font(a, size=8.5, color=MUTED)
    add_field(fp, "PAGE")
    b = fp.add_run(" / ")
    set_run_font(b, size=8.5, color=MUTED)
    add_field(fp, "NUMPAGES")

    doc.core_properties.title = "心语大学生心理健康互助平台项目计划书"
    doc.core_properties.subject = "东南大学成贤学院大学生创新大赛"
    doc.core_properties.author = "心语科创团队"
    doc.core_properties.comments = ""
    return doc


def cover_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("东南大学成贤学院大学生创新大赛")
    set_run_font(r, name=FONT_HEAD, size=12, color=TEAL, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("项目计划书")
    set_run_font(r, name=FONT_HEAD, size=27, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("“心语”大学生心理健康互助平台")
    set_run_font(r, name=FONT_HEAD, size=20, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(33)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("让心理支持更早被看见，让求助更容易发生")
    set_run_font(r, size=11.5, color=MUTED, italic=True)

    meta = doc.add_table(rows=7, cols=2)
    set_table_geometry(meta, [2200, 7160])
    meta_values = [
        ("项目类别", "高教主赛道 · 创意组"),
        ("团队名称", "心语科创团队"),
        ("项目负责人", "杨默涵"),
        ("指导教师", "田静"),
        ("所在学院", "电子与计算机工程学院"),
        ("大创项目编号", "YCX2026189"),
        ("编制日期", "2026-08-24"),
    ]
    for i, (label, value) in enumerate(meta_values):
        add_meta_row(meta, i, label, value)
        set_cell_shading(meta.rows[i].cells[0], LIGHT_BLUE)
        set_cell_shading(meta.rows[i].cells[1], WHITE)
    set_table_style(meta, header=False, body_size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(5)

    add_callout(
        doc,
        "项目定位",
        "面向高校学生的心理健康自助与轻社区平台。首期目标是完成校内可控演示和小范围可行性验证，不把未经授权的真实危机通知、医疗诊断或大规模商业化作为项目能力。",
        fill=LIGHT_TEAL,
        accent=TEAL,
        size=10.0,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(38)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("三条主线")
    set_run_font(r, name=FONT_HEAD, size=9.5, color=GOLD, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("学生可用  ·  安全可解释  ·  团队可执行")
    set_run_font(r, size=11, color=NAVY, bold=True)


def toc_page(doc):
    doc.add_page_break()
    add_kicker(doc, "Project brief")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("内容目录")
    set_run_font(r, name=FONT_HEAD, size=20, color=NAVY, bold=True)
    entries = [
        ("1", "项目概要", 3),
        ("2", "背景、问题与需求", 4),
        ("3", "用户调研与使用场景", 5),
        ("4", "产品方案与核心流程", 6),
        ("5", "技术路线与系统实现", 7),
        ("6", "隐私安全、伦理边界与内容治理", 8),
        ("7", "创新点、竞争分析与项目壁垒", 9),
        ("8", "市场定位与校内试点运营", 10),
        ("9", "服务模式与可持续发展", 11),
        ("10", "团队、组织与实施机制", 12),
        ("11", "进度计划与验收指标", 13),
        ("12", "财务预算、知识产权与预期成果", 14),
        ("13", "风险、社会价值与未来规划", 15),
    ]
    for number, title, page in entries:
        add_toc_entry(doc, number, title, page)
    add_callout(
        doc,
        "编制口径",
        "本计划书按创新创业大赛常见的“问题导向、项目创新、落地实施、团队协作、社会价值”维度组织内容。所有尚未获得校方授权或尚未完成实证评测的内容，均按“计划、演示或待验证”表述。",
        fill=LIGHT_GOLD,
        accent=GOLD,
        size=9.4,
    )
    add_body(doc, "阅读提示：项目首先是一个由学生团队完成的校园公益型创新实践，商业化仅作为后续可持续运维的探索，不以营收规模作为现阶段主要评价指标。", after=0, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=9.2, color=MUTED)


def page_1_overview(doc):
    doc.add_page_break()
    add_section_heading(doc, "1 项目概要")
    add_callout(doc, "一句话定位", "“心语”通过微信小程序连接自我观察、匿名表达、自动初筛、人工复核与支持资源，为高校心理健康服务提供一个低门槛、可解释、可逐步试点的辅助工具。", fill=LIGHT_TEAL, accent=TEAL)
    add_subheading(doc, "1.1 项目背景与目标")
    add_body(doc, "大学生在学业、人际、就业、家庭和生活节奏变化中，可能经历持续的情绪困扰、焦虑、睡眠波动或不愿向熟人表达的时刻。现有高校心理服务以专业咨询和线下活动为主，资源宝贵但服务容量有限；学生是否主动求助，也会受到隐私顾虑、羞耻感和时间成本的影响。")
    add_body(doc, "项目不试图替代心理咨询，也不把算法结果解释成疾病诊断，而是把“先观察一下、先写下来、先得到一条安全而中性的回应”作为入口，再把需要进一步确认的内容交给受控的后台流程。首期目标是完成本校场景下的可用原型、答辩演示闭环和小范围授权试点准备。")
    add_subheading(doc, "1.2 解决方案摘要")
    rows = [
        ("学生端", "今日/自测/树洞/我的等核心入口；完成自测、查看固定规则结果、匿名表达和轻量支持。", "让学生低压力地开始使用"),
        ("分析与分流", "量表由固定规则计分；树洞文本使用规则与模型辅助初筛；模型失败时不默认公开。", "让结果可解释、可回退"),
        ("管理端", "提供内容审核队列、必要的脱敏上下文、处理决定、干预记录和审计日志。", "让人工判断有记录可追溯"),
        ("安全边界", "高风险或隐私优先内容不进入公开流；学校通知、人工处置和真实身份访问均以授权为前提。", "避免把演示能力写成现实承诺"),
    ]
    add_table(doc, ["组成", "主要内容", "直接价值"], rows, [1500, 5320, 2540], body_size=8.9)
    add_subheading(doc, "1.3 当前基础与阶段判断")
    add_body(doc, "项目已完成需求梳理、题库与评分规则整理、FastAPI 后端、数据库模型、学生端小程序、树洞与风险分流、管理后台、演示数据和本地 AI 容错路径等主要开发项，当前从“能否实现”进入“如何稳定验收、如何清楚演示、如何在授权条件下试点”的阶段。")
    add_body(doc, "当前仍需继续完成学生端真机或开发者工具的运行态回归，并进一步核验校内支持资源、社区审核规则和实际部署授权。因此，本计划书将已完成能力、演示能力和待验证事项明确分开。", after=0)


def page_2_background(doc):
    doc.add_page_break()
    add_section_heading(doc, "2 背景、问题与需求")
    add_subheading(doc, "2.1 政策与行业背景")
    add_body(doc, "教育部门持续强调高校心理健康教育、学生支持体系和数字化工具的协同建设。2026 年教育部关于举办中国国际大学生创新大赛的通知，也强调项目要面向经济社会现实需求，重视成果转化、产学研用融合以及真实、健康、合法的参赛材料。对本项目而言，这意味着“使用了人工智能”不是充分理由，必须说明解决什么校园问题、怎么安全落地、团队做了哪些真实工作。")
    add_body(doc, "从服务现场看，心理中心、辅导员和学生工作队伍需要在有限时间内完成宣传教育、个体咨询、危机识别、记录和转介等工作；学生则更需要一个不要求立刻说明全部情况的入口。心语的价值不在于新增一个“测试页面”，而在于尝试把学生端的轻量表达与后台的受控审核流程连接起来。")
    add_subheading(doc, "2.2 主要痛点")
    rows = [
        ("求助启动成本高", "学生不一定愿意以实名、正式预约的方式开始表达，服务端往往只能等待主动求助。", "提供自我观察和匿名表达的前置入口"),
        ("校园语境难判断", "同一句话在同学之间可能是吐槽、玩笑，也可能是持续困扰；单纯关键词拦截容易误报或漏报。", "规则、模型辅助和人工判断分层处理"),
        ("线上工具与校内工作脱节", "测评或倾诉工具可能有前端体验，却没有审核、记录、复核和交接过程。", "用后台队列和审计记录承接必要人工流程"),
        ("隐私与安全需要同时考虑", "越强调匿名，越要讲清数据用途、可访问角色、保存规则和例外情况。", "真实身份、匿名身份、内容审核和演示数据分开治理"),
    ]
    add_table(doc, ["痛点", "具体表现", "项目回应"], rows, [1800, 4620, 2940], body_size=8.8)
    add_subheading(doc, "2.3 用户核心需求")
    for item in [
        "学生：无需下载独立 App，能够用较低门槛完成自测，看到事实性、非诊断式的结果说明。",
        "学生：可以用匿名昵称表达近况，获得“抱抱、陪伴、加油”等不制造社交压力的轻量支持。",
        "审核人员：只在必要权限内查看脱敏内容，知道条目为何进入队列，并能提交明确决定和处理记录。",
        "学校：如果未来获得授权，可以在既有心理工作流程中使用该工具，而不是被迫接收一个无法解释的黑箱系统。",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "边界先行", "平台不是医疗诊断系统，不输出疾病诊断；AI 不作最终审核或危机处置决定；没有正式授权、值班角色和响应机制时，不承诺学校会自动联系学生。", fill=LIGHT_GOLD, accent=GOLD, size=9.5)


def page_3_research(doc):
    doc.add_page_break()
    add_section_heading(doc, "3 用户调研与使用场景")
    add_subheading(doc, "3.1 调研基础与证据边界")
    add_body(doc, "前期工作已覆盖文献查阅、需求梳理、校园心理文本与题库资料整理、产品流程讨论以及学生端和管理端原型实现。项目现有材料能够支持“问题存在、方案可实现、闭环可演示”的判断，但尚未形成可代表全校学生的严格抽样统计，也没有获得可公开展示的真实危机个案数据。因此，本任务书不虚构用户比例、模型准确率或试点转化率，而采用可复核的场景假设来指导下一阶段验证。")
    add_subheading(doc, "3.2 典型用户画像")
    rows = [
        ("A 类：想先自己看看", "近期有压力或睡眠波动，但还不确定是否需要正式求助。", "进入自测，阅读固定结果；需要时查看支持资源。"),
        ("B 类：想找个地方写下来", "不希望被熟人立刻认出，想表达一件学习、人际或生活中的小事。", "选择匿名身份、选话题、发布文字；内容按规则进入公开或审核流程。"),
        ("C 类：需要处理内容的人", "心理中心、辅导员或受授权的审核角色，需要在有限时间内判断内容状态。", "在后台查看脱敏内容、处理决定、交接备注和审计记录。"),
    ]
    add_table(doc, ["用户", "状态与顾虑", "希望完成的任务"], rows, [1660, 3900, 3800], body_size=8.9)
    add_subheading(doc, "3.3 三类高频使用场景")
    add_numbered(doc, "自我观察场景：学生从测评入口开始，按“一题一页”完成自测，系统按版本化固定规则生成分数或定性观察，结果页先说明事实，再给出自主支持入口。")
    add_numbered(doc, "匿名表达场景：学生进入树洞，首次发帖前单独了解社区审核规则，输入 1 至 800 个非空字符并选择话题；平台不支持图片、链接、私信和自由文本评论。")
    add_numbered(doc, "人工复核场景：内容经过自动检查后，低风险内容可进入公开流；敏感、模糊、隐私或安全优先内容进入等待确认、保护展示或受限流程，由人工作出最终决定。")
    add_subheading(doc, "3.4 下一阶段验证方式")
    add_body(doc, "在获得校内授权后，团队计划以志愿学生、小规模、可撤回的方式开展可用性访谈和流程演练，重点观察：学生是否理解匿名与审核边界、结果文案是否容易被误读、发帖流程是否足够轻量、后台人员是否能在不接触多余个人信息的情况下完成判断。若授权条件尚未具备，则继续使用合成资料完成开发测试和答辩演示。", after=0)


def page_4_product(doc):
    doc.add_page_break()
    add_section_heading(doc, "4 产品方案与核心流程")
    add_callout(doc, "产品原则", "少一点功能堆叠，多一条可以解释的服务链路：先让学生愿意开始，再让系统把需要确认的内容交给合适的人。", fill=LIGHT_BLUE, accent=BLUE)
    add_subheading(doc, "4.1 产品组成")
    rows = [
        ("今日 / 自测", "每日入口、情绪/焦虑/睡眠等自我观察、结果与历史回看。", "固定规则优先，不用风险标签制造压力。"),
        ("树洞轻社区", "匿名昵称、话题、帖子详情、回应与轻量支持。", "不开放私信、图片、链接或热度排行。"),
        ("后台工作区", "内容审核、预警/跟进、帖子管理、用户目录、审计日志和统计。", "按角色授权，只展示判断所需信息。"),
        ("演示与容错", "演示账号、合成案例、本地 mock AI、答辩 runbook。", "明确标记“演示模式”，不伪造真实学校处置。"),
    ]
    add_table(doc, ["模块", "功能范围", "设计边界"], rows, [1700, 4940, 2720], body_size=8.8)
    add_subheading(doc, "4.2 核心闭环")
    flow = [
        ("1", "学生输入", "自测或匿名表达"),
        ("2", "固定规则", "计分、边界和安全文案"),
        ("3", "辅助初筛", "规则 + AI，失败不默认公开"),
        ("4", "人工判断", "公开、保护、暂缓、驳回或升级"),
        ("5", "支持与记录", "资源入口、处理记录、审计"),
    ]
    add_table(doc, ["阶段", "动作", "结果"], flow, [1000, 2300, 6060], header_fill=LIGHT_TEAL, body_size=8.8)
    add_subheading(doc, "4.3 学生端体验要点")
    for item in [
        "测评答题过程中隐藏底部导航，避免中途被其他功能打断；结果页先显示固定规则生成的一句话总结，再展示分数、定位或定性观察。",
        "树洞浏览不强制先同意发帖审核；首次写下帖子或回应时，再展示独立的社区审核说明。",
        "内容公开后只显示匿名身份快照；支持反馈可以切换或撤回，但不显示支持数量、支持者名单或热度排名。",
        "进入安全支持状态时，优先呈现经授权的支持资源；不显示模型置信度、内部命中词或“高危/低危”等后台标签。",
    ]:
        add_bullet(doc, item, size=9.7)


def page_5_technology(doc):
    doc.add_page_break()
    add_section_heading(doc, "5 技术路线与系统实现")
    add_subheading(doc, "5.1 技术架构")
    rows = [
        ("学生端", "微信原生小程序 / JavaScript / WXML / WXSS", "承载登录、同意、自测、报告、树洞和帮助资源。"),
        ("服务端", "FastAPI + Pydantic + SQLAlchemy + Alembic", "提供鉴权、问卷、评分、树洞、风险分流、后台接口。"),
        ("管理端", "Streamlit 网页后台", "面向电脑浏览器，处理队列、帖子、审计和统计。"),
        ("数据层", "MySQL 目标基线；演示环境支持本地轻量数据库路径", "保存题库、答案快照、帖子状态、工单和审计记录。"),
        ("AI 接口", "可替换的大模型客户端 + 本地 mock fallback", "仅对树洞文本做辅助分析，不接收真实身份和完整测评答案。"),
    ]
    add_table(doc, ["层级", "选型", "承担职责"], rows, [1500, 3300, 4660], body_size=8.7)
    add_subheading(doc, "5.2 评分与风险分流逻辑")
    add_body(doc, "量表评分由题库中的选项映射、反向计分和版本化阈值完成，AI 不参与分数计算，不生成医疗结论。树洞文本则经过内容合规与风险辅助分析：低风险或需要关注但可公开的内容进入相应公开状态；高风险、隐私优先、模糊或自动流程失败的内容不默认公开，并进入人工确认或安全支持路径。历史测评结果可以形成后台跟进提示，但不直接把一条当前低风险文字拦截掉。")
    add_subheading(doc, "5.3 当前开发基础")
    for item in [
        "已完成 FastAPI 应用骨架、环境配置、数据库模型、迁移链路和核心账户/授权/审计数据域。",
        "已完成五类题库种子、评分服务、反向计分、硬触发规则、单量表结果和完整报告解锁逻辑；完整报告由四份必做问卷完成状态控制，辅助题组不阻塞解锁。",
        "已完成学生端问卷、报告、树洞、高风险拦截、删除与轻量支持；已完成管理员登录、队列、帖子管理、审计和统计页面。",
        "已完成演示数据服务、mock AI 和关键后端测试；学生端真实渲染、真机交互和校内支持资源仍需在后续验收中核对。",
    ]:
        add_bullet(doc, item, size=9.5)
    add_callout(doc, "技术承诺边界", "项目当前不宣称“校园心理文本微调模型已达到某一固定准确率”。后续如获得合法、脱敏、可用于研究的数据，才开展独立评测；在此之前，以流程可靠性、可回退和人工复核来降低单一模型的不确定性。", fill=LIGHT_GOLD, accent=GOLD, size=9.3)


def page_6_safety(doc):
    doc.add_page_break()
    add_section_heading(doc, "6 隐私安全、伦理边界与内容治理")
    add_subheading(doc, "6.1 数据与同意")
    add_body(doc, "学生端的基础服务说明、身份核验说明和社区审核说明分层呈现；浏览树洞不等于同意发帖审核。真实身份资料与社区匿名资料分开存储、分开授权和分开查询，公开页面只呈现系统生成的匿名昵称与抽象头像。演示模式只使用演示账号和合成资料，不要求填写或展示真实学生身份。")
    add_body(doc, "平台不向大模型发送真实姓名、学号、手机号或完整测评答案；模型输入在必要时只保留经脱敏的树洞文本。任何“匿名”表述都以实际权限、保存规则和后台审计为准，不使用“绝不会泄露”等无法绝对保证的承诺。")
    add_subheading(doc, "6.2 控制措施")
    rows = [
        ("前台匿名", "社区公开身份与真实身份分离；历史内容保留发布时的匿名快照。", "降低被同学直接识别的可能"),
        ("后台最小可见", "队列表格不显示正文、真实姓名、学号、心理分数或作者历史；详情只展示必要脱敏信息。", "减少越权浏览"),
        ("权限与审计", "管理员登录、角色控制、敏感详情访问记录、决定与处理时间进入审计日志。", "让关键动作可追踪"),
        ("软删除与演示隔离", "学生删除使用软删除语义；演示案例与真实账号分离，界面持续标记演示状态。", "兼顾治理需要与演示可信度"),
    ]
    add_table(doc, ["控制点", "实现方向", "目的"], rows, [1500, 5000, 2960], header_fill=LIGHT_BLUE, body_size=8.7)
    add_subheading(doc, "6.3 内容治理状态")
    add_numbered(doc, "公开：内容通过自动检查，未发现需要进一步确认的情况。")
    add_numbered(doc, "保护方式公开：内容对读者先折叠，用户主动展开后再阅读，避免未经准备的围观。")
    add_numbered(doc, "暂不公开：自动检查未完成、含个人信息或判断不充分，等待人工确认。")
    add_numbered(doc, "驳回或升级复核：违反社区规则或出现安全线索时，由受授权人员作出处理；当前演示只记录模拟流程，不自动向真实学校或家人发送通知。")
    add_callout(doc, "伦理底线", "AI 只提供辅助提示；最终公开、驳回、保护或升级由人工决定。平台不提供诊断、治疗、危机处置承诺，也不将学生的心理状态变成公开标签。", fill=LIGHT_GOLD, accent=GOLD, size=9.4)


def page_7_innovation(doc):
    doc.add_page_break()
    add_section_heading(doc, "7 创新点、竞争分析与项目壁垒")
    add_subheading(doc, "7.1 四层创新")
    rows = [
        ("产品创新", "把自我观察、匿名表达和固定支持入口放在同一条低压力使用路径上。", "学生不必一开始就进入正式咨询，也不把树洞做成无限制论坛。"),
        ("流程创新", "从“学生输入”到“自动检查”再到“人工复核”，将需要确认的内容转化为可操作的队列。", "把情绪表达和学校心理工作之间的断点变成流程连接。"),
        ("治理创新", "固定规则负责计分和安全文案，AI 只作辅助，后台决定和审计记录形成闭环。", "降低黑箱判断和自动化误处置风险。"),
        ("场景创新", "围绕高校学生的学业、人际、睡眠和就业压力设计话题、文案和权限。", "不把通用互联网社区的热度机制直接搬入校园心理场景。"),
    ]
    add_table(doc, ["创新层", "具体做法", "相对价值"], rows, [1500, 4920, 3040], header_fill=LIGHT_TEAL, body_size=8.7)
    add_subheading(doc, "7.2 竞品与替代方案对比")
    rows = [
        ("单一心理测评工具", "通常聚焦问卷与结果展示", "心语增加匿名表达、内容治理和后台审核闭环"),
        ("商业心理咨询应用", "以付费咨询或内容服务为主", "心语首期面向校内学生免费，以学校现有支持体系为承接"),
        ("开放式匿名社区", "表达门槛低，但内容治理和校园边界不一定清晰", "心语限制互动形态，强调匿名、脱敏、人工复核与审计"),
        ("纯 AI 聊天或情绪分析", "反馈即时，但容易被误解为专业判断", "心语把 AI 放在辅助位置，固定规则和人工判断优先"),
    ]
    add_table(doc, ["替代形态", "常见优势", "心语的差异"], rows, [2000, 3160, 4300], body_size=8.7)
    add_subheading(doc, "7.3 可积累的项目壁垒")
    for item in [
        "围绕校园心理服务形成的流程经验：同意、匿名、发布、复核、处置记录之间的边界比单个页面更难复制。",
        "可版本化的题库、固定结果文案、风险规则和后台审计结构，便于在不同学校授权后调整，而不是每次从零开始。",
        "学生团队与指导教师共同完成产品、技术、测试和文档，能够把课程知识转化为可运行的校园场景原型。",
    ]:
        add_bullet(doc, item, size=9.6)
    add_callout(doc, "克制的竞争判断", "项目目前的优势是“流程完整、边界清楚、可在校内验证”，不是模型性能领先或商业规模领先。真正的壁垒需要通过授权试点、用户反馈和长期治理逐步形成。", fill=LIGHT_GOLD, accent=GOLD, size=9.4)


def page_8_market(doc):
    doc.add_page_break()
    add_section_heading(doc, "8 市场定位与校内试点运营")
    add_subheading(doc, "8.1 目标对象")
    rows = [
        ("第一层", "东南大学成贤学院在校学生", "先验证学生是否愿意使用、流程是否清楚、内容治理是否可执行。"),
        ("第二层", "具有类似心理健康教育需求的本地高校", "仅在校内试点获得授权、数据和运维经验后再考虑复制。"),
        ("合作对象", "学生处、心理健康教育中心、心理工作站、团学组织与信息化支持部门", "提供资源审核、流程授权、用户招募和场景反馈。"),
    ]
    add_table(doc, ["层级", "对象", "当前任务"], rows, [1500, 4200, 3760], body_size=8.8)
    add_subheading(doc, "8.2 三阶段试点路径")
    add_numbered(doc, "阶段一：内部可演示。使用合成账号和案例，完成登录、测评、报告、树洞、拦截、人工审核和审计链路的稳定彩排。")
    add_numbered(doc, "阶段二：授权小范围验证。由学校相关部门确认支持资源、审核角色和数据规则后，邀请志愿学生进行可撤回的可用性体验和访谈，不追求短期用户数量。")
    add_numbered(doc, "阶段三：形成可复制包。沉淀题库配置、审核规则、权限说明、培训材料、部署文档和反馈表；只有在前一阶段问题得到解决后，才讨论向相近高校推广。")
    add_subheading(doc, "8.3 低成本运营方式")
    for item in [
        "依托学校心理工作站、心理协会、班级主题活动和官方公众号进行说明，不采用夸张的焦虑营销或付费投放。",
        "把“如何使用、数据怎么处理、哪些情况不会公开”作为推广的主要内容，让学生在知情基础上选择是否使用。",
        "通过匿名反馈、可用性访谈、后台审核复盘和故障记录迭代产品，优先修复误解和安全问题，再扩展功能。",
        "建立“资源有效性检查”清单，支持电话、服务时间和校内入口在正式使用前必须由授权部门确认。",
    ]:
        add_bullet(doc, item, size=9.6)
    add_callout(doc, "试点原则", "先小范围、后复制；先确认授权和支持资源、后公开推广；先验证学生体验和审核流程、后讨论规模化。", fill=LIGHT_TEAL, accent=TEAL, size=9.5)


def page_9_service(doc):
    doc.add_page_break()
    add_section_heading(doc, "9 服务模式与可持续发展")
    add_subheading(doc, "9.1 公益属性优先")
    add_body(doc, "心语首期服务对象是本校学生，核心自测、树洞和支持资源不向学生收取费用，不通过出售敏感数据、广告植入或心理结果分层获利。项目更适合被理解为“高校心理健康教育的数字化辅助工具”，而不是面向个人用户的高价心理服务产品。")
    add_subheading(doc, "9.2 现实可行的支持来源")
    rows = [
        ("当前", "学校大创项目经费、指导教师与团队成员的时间投入、学校已有设备和开发环境。", "完成原型、测试和答辩材料。"),
        ("试点期", "学校相关部门的场景授权、心理专业支持、信息化部署协助和小规模试点资源。", "验证流程与内容，而不是追求营收。"),
        ("后续探索", "在合法合规、需求明确的前提下，讨论部署配置、培训和运维等低成本技术服务，或争取学校专项经费、公益项目资助。", "只覆盖持续维护成本，不以扩大数据采集为目标。"),
    ]
    add_table(doc, ["阶段", "资源来源", "主要用途"], rows, [1500, 5320, 2640], header_fill=LIGHT_BLUE, body_size=8.7)
    add_subheading(doc, "9.3 服务交付边界")
    for item in [
        "学生端：提供清晰、低门槛的自我观察和表达工具，不提供诊断、治疗或强制性结论。",
        "学校端：提供可配置的题库、审核角色、队列和审计能力；真实部署前需要确认数据处理、值班与处置协议。",
        "团队端：形成安装说明、演示账号、测试用例、人工演练脚本和问题清单，保证团队成员更替后仍可继续维护。",
    ]:
        add_bullet(doc, item, size=9.7)
    add_subheading(doc, "9.4 可持续性的判断标准")
    add_body(doc, "项目能否持续，不只取决于是否收费，还取决于学校是否愿意使用、学生是否信任、审核人员是否能完成工作、题库和支持资源是否有人维护。未来评价项目时，将优先看有效使用、反馈质量、审核闭环和安全记录，而不是简单看注册人数或收入。", after=0)


def page_10_team(doc):
    doc.add_page_break()
    add_section_heading(doc, "10 团队、组织与实施机制")
    add_subheading(doc, "10.1 团队成员与分工")
    rows = [
        ("杨默涵", "第一主持人", "产品方案、后端开发、项目统筹", "把需求、技术和答辩材料串成完整闭环"),
        ("李晓曼", "第二主持人", "小程序前端、用户调研、运营策划", "负责学生端体验与试点反馈整理"),
        ("王启轩", "项目成员", "题库/数据整理、分析规则与模型调试", "保证评分与辅助分析逻辑可解释"),
        ("王项玉", "项目成员", "后台管理、测试与文档整理", "保证审核、审计和演示材料完整"),
        ("田静", "指导教师", "技术指导、方案审核、进度把关", "在产品边界、研究规范和成果表达上提供指导"),
    ]
    add_table(doc, ["成员", "身份", "主要分工", "对项目的直接贡献"], rows, [1500, 1400, 3060, 3400], body_size=8.4)
    add_subheading(doc, "10.2 组织方式")
    add_body(doc, "团队采用“产品与用户、技术与数据、后台与测试、教师指导”四类角色协作。每周以任务清单和演示结果同步进度；涉及风险规则、隐私文案、真实资源和对外试点的变更，先由团队内部完成影响评估，再提交指导教师和相关专业人员审核。")
    add_subheading(doc, "10.3 团队优势与能力补齐")
    rows = [
        ("已有优势", "成员来自电子与计算机工程学院，具备小程序、后端、数据库、测试和 AI 接口开发基础；项目已经从文档进入可运行原型。"),
        ("当前短板", "心理测量、危机干预、隐私合规和真实校园工作流程需要专业人员与校方共同参与，不能由学生团队单独推断。"),
        ("补齐方式", "把专业审核、支持资源核验和授权试点列为下一阶段前置条件；团队负责工程实现、文档和可复核测试，不替代心理专业判断。"),
    ]
    add_table(doc, ["方面", "说明"], rows, [1700, 7660], header_fill=LIGHT_TEAL, body_size=8.8)
    add_callout(doc, "团队承诺", "不把未完成的专业合作、模型评测或校内试点写成既成成果；所有阶段目标都以“能否验证、谁来负责、出现问题如何回退”为验收标准。", fill=LIGHT_GOLD, accent=GOLD, size=9.4)


def page_11_schedule(doc):
    doc.add_page_break()
    add_section_heading(doc, "11 进度计划与验收指标")
    add_subheading(doc, "11.1 项目进度")
    rows = [
        ("2025-12 至 2026-02", "已完成", "文献查阅、需求调研、问卷与校园语境资料整理，形成初版产品和技术方案。"),
        ("2026-03 至 2026-06", "已完成", "完成后端基础设施、数据模型、题库导入、评分与报告、学生端问卷及报告流程。"),
        ("2026-07 至 2026-08", "已完成/收口", "完成树洞风险分流、后台审核、审计、演示数据和前端频道化调整，进入验收与答辩准备。"),
        ("2026-09 至 2026-10", "计划", "完成学生端运行态回归、可用性检查、演示脚本和问题修复；核对支持资源与边界文案。"),
        ("2026-11 至 2026-12", "计划", "在获得授权时准备小范围校内试点；完成研究总结、软著材料、结题材料和后续迭代清单。"),
    ]
    add_table(doc, ["时间", "状态", "主要工作"], rows, [2100, 1500, 5760], body_size=8.7)
    add_subheading(doc, "11.2 验收指标与大赛评价映射")
    rows = [
        ("问题导向", "能说明学生为何需要低门槛入口，以及学校为何需要受控的审核流程。", "需求材料、用户场景、问题-方案映射"),
        ("项目创新", "自我观察、轻社区、AI 辅助、人审与审计形成一条校园场景闭环。", "可运行原型、流程图、对比分析"),
        ("落地实施", "学生端、服务端、后台端和演示容错均可运行；授权试点按条件推进。", "代码、演示 runbook、联调记录"),
        ("团队协作", "成员分工真实，产品、开发、测试、文档和教师指导形成协作机制。", "任务分工、进度记录、会议与测试材料"),
        ("社会价值", "学生免费、公益属性优先，提升自我观察和校园心理支持的可达性。", "边界声明、试点方案、反馈与社会价值分析"),
    ]
    add_table(doc, ["评价维度", "本项目的回答", "验收证据"], rows, [1800, 4900, 2660], header_fill=LIGHT_TEAL, body_size=8.5)
    add_callout(doc, "阶段性结论", "项目功能面已基本完成，下一阶段的关键不是继续堆叠功能，而是把运行态稳定性、专业边界、用户理解和演示可信度逐项验收。", fill=LIGHT_BLUE, accent=BLUE, size=9.4)


def page_12_budget(doc):
    doc.add_page_break()
    add_section_heading(doc, "12 财务预算、知识产权与预期成果")
    add_subheading(doc, "12.1 经费预算")
    add_body(doc, "项目现有计划总经费为 1000 元，全部来自学校大创项目资助，自筹 0 元。预算以完成软件原型、资料整理、测试调试和成果申报为目的，不按商业项目配置服务器、广告投放或大规模采购。")
    rows = [
        ("业务费", "软件著作权申请费", "700", "用于成果申报及相关材料准备"),
        ("业务费", "文献检索与资料整理", "100", "文献查阅、调研资料打印和归档"),
        ("仪器设备购置费", "外设配件与调试耗材", "100", "用于本地开发、联调与演示"),
        ("材料费", "打印与宣传物料", "100", "用于展示、说明材料和小范围活动"),
        ("合计", "", "1000", "学校大创资助；自筹 0 元"),
    ]
    add_table(doc, ["科目", "具体支出", "金额（元）", "用途说明"], rows, [1800, 3000, 1300, 3260], header_fill=LIGHT_GOLD, body_size=8.6, alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])
    add_subheading(doc, "12.2 知识产权与成果")
    for item in [
        "软件成果：心语大学生心理健康互助微信小程序、后端 API 服务、管理后台和可演示的风险复核闭环。",
        "数据与规则成果：版本化题库、评分映射、结果文案、树洞审核状态、测试用例和演示案例说明。",
        "文档成果：需求与流程文档、技术说明、调研/测试记录、人工演练脚本、用户反馈表和项目研究报告。",
        "知识产权：计划申请“心语”大学生心理健康互助平台 V1.0 软件著作权，具体权利归属与申报内容按学校要求办理。",
    ]:
        add_bullet(doc, item, size=9.5)
    add_subheading(doc, "12.3 预期成果的层次")
    rows = [
        ("可交付", "能启动、能演示、能测试的学生端和管理端原型；关键操作有状态反馈和错误回退。"),
        ("可复核", "评分、发布、审核、审计和演示模式有明确规则与测试材料，评委可以看懂项目如何工作。"),
        ("可延续", "在授权与专业支持条件具备时，可进一步完成校内小规模验证；不将尚未验证的外部推广写成既定成果。"),
    ]
    add_table(doc, ["层次", "成果描述"], rows, [1700, 7660], header_fill=LIGHT_TEAL, body_size=8.8)


def page_13_risk_social(doc):
    doc.add_page_break()
    add_section_heading(doc, "13 风险、社会价值与未来规划")
    add_subheading(doc, "13.1 风险识别与应对")
    rows = [
        ("模型误判", "误报或漏报导致内容状态不合适", "AI 只作辅助；失败不默认公开；人工复核；持续记录样本问题", "团队技术负责人"),
        ("隐私泄露", "学生不再信任平台，产生合规风险", "最小化采集、匿名与真实身份分离、权限控制、脱敏与审计", "后端/后台负责人"),
        ("过度承诺", "被误解为诊断或自动危机干预系统", "所有页面与材料明确“不构成诊断”；学校通知与支持资源以授权为前提", "指导教师/产品负责人"),
        ("用户不愿使用", "试点无法获得有效反馈", "先做可用性访谈，解释审核边界，不用焦虑营销；允许退出和撤回", "前端/运营负责人"),
        ("校内流程难落地", "没有明确审核角色、值班规则和支持资源", "先演示、后授权；将流程、联系人、时间和责任写成书面规则", "项目负责人"),
    ]
    add_table(doc, ["风险", "可能影响", "应对措施", "责任"], rows, [1500, 2420, 4560, 880], header_fill=LIGHT_GOLD, body_size=8.1)
    add_subheading(doc, "13.2 社会价值")
    for item in [
        "对学生：提供免费、低门槛、非诊断式的自我观察和匿名表达入口，帮助学生在正式求助之前先看见自己的状态。",
        "对学校：将学生端的内容治理、人工复核、审计记录和支持资源整理成可讨论的流程原型，为心理健康教育数字化提供实践案例。",
        "对团队：让计算机专业学生在真实社会问题中学习需求调研、数据治理、软件工程、人工智能边界和跨专业协作。",
    ]:
        add_bullet(doc, item, size=9.5)
    add_subheading(doc, "13.3 未来规划")
    rows = [
        ("近期：6-12 个月", "完成真机/开发者工具回归、文案与资源核验、软件著作权申报和答辩材料；条件具备时开展授权小范围体验。"),
        ("中期：1-2 年", "根据试点反馈优化题库、结果表达、审核工作台和数据审计；建立跨角色培训与交接材料。"),
        ("长期：视条件而定", "在获得学校、专业人员和数据治理支持后，评估向相近高校复制的可行性；商业化只作为运维可持续性的辅助探索。"),
    ]
    add_table(doc, ["阶段", "规划"], rows, [2200, 7460], header_fill=LIGHT_TEAL, body_size=8.7)
    add_callout(doc, "结语", "心语最先要证明的不是能服务多少人，而是能否在校园场景里把一条小而完整的心理支持流程做得清楚、可靠、可复核。", fill=LIGHT_BLUE, accent=BLUE, size=9.6)
    add_minor_heading(doc, "参考依据与项目材料")
    add_body(doc, "1. 教育部：《教育部关于举办中国国际大学生创新大赛（2026）的通知》，2026。\n2. 中国国际大学生创新大赛（2025）评审规则，公开版。\n3. 心语 V2 已确认产品需求与决策记录、评估医学依据与参考、学生端/后台设计方案。\n4. 心语项目 PRD、应用流程、技术栈、实施计划、进度摘要与测试/演练材料。", after=0, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=8.1, color=MUTED)


def build():
    doc = setup_document()
    cover_page(doc)
    toc_page(doc)
    page_1_overview(doc)
    page_2_background(doc)
    page_3_research(doc)
    page_4_product(doc)
    page_5_technology(doc)
    page_6_safety(doc)
    page_7_innovation(doc)
    page_8_market(doc)
    page_9_service(doc)
    page_10_team(doc)
    page_11_schedule(doc)
    page_12_budget(doc)
    page_13_risk_social(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
